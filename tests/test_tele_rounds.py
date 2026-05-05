from __future__ import annotations

from io import BytesIO
from unittest.mock import patch

import pytest

from src.tele_rounds import parse_docx_patient_blocks, process_icu_report

try:
    from docx import Document
except Exception:
    Document = None

SAMPLE_TEXT = """
ICU RMO REPORTING FORMAT
SECTION 1: PATIENT IDENTIFICATION & BASIC DETAILS
UHID Patient Name Age/Sex Bed DOA Primary Diagnosis
202601170031 RAM KUMAR 75/M 14 08/04/26 DR SB MISHRA CKD ON MHD, POST CPR STATUS ICU
SECTION 2: CLINICAL STATUS SUMMARY (CURRENT)
29.04.26 7.30 AM E1VTM1 107 161/96(106)
SECTION 9: INVESTIGATIONS & INTERVENTIONS
29.04.26 7.30 AM Y EEG REPORT PENDING TRACHEOSTOMY DONE HD TODAY
SECTION 10: RED FLAG IDENTIFICATION (CRITICAL)
29.04.26 7.30 AM Y LOW GCS Y
SECTION 11: SUMMARY FOR CONSULTANT (MOST IMPORTANT)
29.04.26 7.30 AM CRITICAL MAJOR CONCERN SHOCK WORSENED REQUIREMENT VENT SUPPORT
SECTION 12: ESCALATION & ORDERS
29.04.26 7.30 AM DR SATYAJIT REPEAT ABG AND CONTINUE NORAD YES
SECTION 13: HANDOVER ACCOUNTABILITY
29.04.26 7.30 AM RMO A
"""


def test_process_icu_report_docx_uses_section_blocks() -> None:
    payload = {"raw_text": SAMPLE_TEXT, "table_rows": []}
    with patch("src.tele_rounds.extract_text", return_value=payload):
        report = process_icu_report("CCM PT DATA UPDATE ON 22ND FEB EVENING.docx", b"fake")

    patients = report.get("patients", [])
    assert len(patients) == 1
    patient = patients[0]
    assert patient.get("bed") == "14"
    assert patient.get("red_flag") is True
    assert "SHOCK" in str(patient.get("major_concern", "")).upper()
    assert "VENT SUPPORT" in str(patient.get("rmo_recommendation", "")).upper()


def test_process_icu_report_unsupported_file_type() -> None:
    try:
        process_icu_report("rounds.txt", b"hello")
    except Exception as error:  # noqa: BLE001
        assert "Only PDF and DOCX" in str(error)
    else:
        raise AssertionError("Expected unsupported file type error")


def _add_table(document: object, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value


@pytest.mark.skipif(Document is None, reason="python-docx not installed")
def test_parse_docx_patient_blocks_maps_latest_first_data_rows() -> None:
    document = Document()
    _add_table(
        document,
        [
            ["SECTION 1: PATIENT IDENTIFICATION", "", "", "", ""],
            ["UHID", "Patient Name", "Bed No", "Category", "Primary Diagnosis"],
            ["P001", "A Patient", "17", "ICU", "Septic shock"],
        ],
    )
    _add_table(
        document,
        [
            ["SECTION 2: CLINICAL STATUS SUMMARY", "", "", "", ""],
            ["Sl No", "GCS", "BP (MAP)", "SpO2", "RR"],
            ["05.05.26 7.30 AM", "E4V5M6", "156/70(85)", "94", "24"],
            ["04.05.26 9.30 PM", "E3V4M5", "100/60(70)", "90", "30"],
        ],
    )
    _add_table(
        document,
        [
            ["SECTION 10: RED FLAG IDENTIFICATION", "", ""],
            ["Sl No", "Red Flag Present", "Reason"],
            ["05.05.26", "Y", "Hypoxia"],
        ],
    )
    _add_table(
        document,
        [
            ["SECTION 11: SUMMARY FOR CONSULTANT", "", "", ""],
            ["Sl No", "Current Status", "Major Concern", "RMO Recommendation"],
            ["05.05.26", "CRITICAL", "Septic shock on support", "Review source control"],
        ],
    )
    _add_table(
        document,
        [
            ["SECTION 12: ESCALATION & ORDERS", ""],
            ["Sl No", "Advice Given"],
            ["05.05.26", "Continue norad and repeat lactate"],
        ],
    )
    _add_table(
        document,
        [
            ["SECTION 13: HANDOVER ACCOUNTABILITY", ""],
            ["Sl No", "RMO"],
            ["05.05.26", "RMO A"],
        ],
    )

    buffer = BytesIO()
    document.save(buffer)

    rows, warnings = parse_docx_patient_blocks(buffer.getvalue())
    assert warnings == []
    assert len(rows) == 1
    patient = rows[0]
    assert patient["bed"] == "17"
    assert patient["patient_id"] == "P001"
    assert patient["patient_name"] == "A Patient"
    assert patient["category"] == "ICU"
    assert patient["gcs"] == "E4V5M6"
    assert patient["bp_map"] == "156/70(85)"
    assert patient["spo2"] == "94"
    assert patient["rr"] == "24"
    assert "100/60" not in patient["section2_status"]
    assert patient["red_flag_present"] == "Y"
    assert patient["major_concern"] == "Septic shock on support"
    assert patient["rmo_recommendation"] == "Review source control"
    assert patient["section12_orders"] == "Continue norad and repeat lactate"
