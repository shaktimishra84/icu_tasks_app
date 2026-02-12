from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass
from datetime import datetime
from functools import lru_cache
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception:
    Document = None

try:
    import yaml
except Exception:
    yaml = None


REPO_ROOT = Path(__file__).resolve().parents[1]
RULES_PATH = REPO_ROOT / "rules" / "rules_v1.yaml"
ALGORITHM_ROOT = REPO_ROOT / "rules" / "algorithms"

DISPLAY_COLUMNS = [
    "Bed",
    "Patient ID",
    "Diagnosis",
    "Status",
    "Supports",
    "System tag",
    "Matched algorithms",
    "Missing Tests",
    "Missing Imaging",
    "Missing Consults",
    "Care checks (deterministic)",
    "Pending (verbatim)",
    "Key labs/imaging (1 line)",
    "Round trend",
    "Deterioration since last round",
    "Deterioration reasons",
]

SYNONYM_GROUPS = [
    {"eeg", "electroencephalogram"},
    {"2d echo", "echo", "echocardiography", "echo report"},
    {"ctpa", "ct pulmonary angiography"},
    {"nephrology", "nephro", "nephrology call"},
    {"cardiology", "cardio", "cardiology call"},
    {"gastro", "gi", "gastro call", "gastroenterology"},
    {"neurology", "neuro", "neuro consult"},
    {"ugie", "endoscopy", "upper gi endoscopy"},
    {"prbc", "blood products", "transfusion"},
    {"vap bundle", "ventilator bundle"},
    {"ventilator review", "vent settings", "ventilator settings", "vent review"},
]

OLDER_FILENAME_HINTS = [
    "old",
    "older",
    "previous",
    "prev",
    "last",
    "morning",
    "_am",
    "-am",
    "round1",
    "r1",
]

NEWER_FILENAME_HINTS = [
    "new",
    "newer",
    "latest",
    "current",
    "today",
    "updated",
    "evening",
    "_pm",
    "-pm",
    "round2",
    "r2",
]


@dataclass
class RequiredItem:
    category: str
    name: str
    priority: str
    trigger_keywords: list[str]
    evidence_keywords: list[str]


@dataclass
class AlgorithmSpec:
    name: str
    system_tag: str
    diagnosis_keywords: list[str]
    supports_keywords: list[str]
    lab_keywords: list[str]
    strict_diagnosis_keywords: list[str]
    required_items: list[RequiredItem]


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _normalized_key(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", text.lower())


def _contains_any(text: str, needles: list[str]) -> bool:
    hay = text.lower()
    return any(needle.lower() in hay for needle in needles)


def _contains_strict_keyword(text: str, keyword: str) -> bool:
    pattern = r"\b" + re.escape(keyword.lower()) + r"\b"
    return re.search(pattern, text.lower()) is not None


def _priority_rank(priority: str) -> int:
    return {"high": 0, "medium": 1, "low": 2}.get(priority.lower(), 3)


def _status_group(status: str) -> str:
    lower = status.lower()
    if "deceased" in lower or "dead" in lower:
        return "DECEASED"
    if "critical" in lower or "crit" in lower:
        return "CRITICAL"
    if "sick" in lower:
        return "SICK"
    if "serious" in lower:
        return "SERIOUS"
    return "OTHER"


def _status_severity(status_group: str) -> int:
    return {
        "OTHER": 0,
        "SERIOUS": 1,
        "SICK": 2,
        "CRITICAL": 3,
        "DECEASED": 4,
    }.get(status_group, 0)


def _dedupe_strings(values: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        key = value.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(value)
    return out


def _patient_key(row: dict[str, Any]) -> str:
    patient_id_key = _normalized_key(str(row.get("Patient ID", "")))
    if patient_id_key:
        return f"pid:{patient_id_key}"
    bed_key = _normalized_key(str(row.get("Bed", "")))
    if bed_key:
        return f"bed:{bed_key}"
    return ""


def _support_burden_score(row: dict[str, Any]) -> float:
    score = 0.0
    if row.get("_is_mv"):
        score += 2.0
    if row.get("_is_niv"):
        score += 1.0
    if row.get("_is_vaso"):
        score += 2.0
    if row.get("_is_rrt"):
        score += 2.0
    if row.get("_is_o2"):
        score += 0.5
    return score


def _support_markers(row: dict[str, str]) -> dict[str, bool]:
    supports_text = " ".join(
        [
            row.get("supports", ""),
            row.get("actions_done", ""),
            row.get("plan_next_12h", ""),
            row.get("pending", ""),
        ]
    ).lower()
    return {
        "mv": _contains_any(supports_text, ["mv", "ventilator", "mechanical ventilation"]),
        "niv": _contains_any(supports_text, ["niv", "bipap", "bi pap", "cpap"]),
        "vaso": _contains_any(
            supports_text,
            ["norad", "norepinephrine", "adrenaline", "epinephrine", "vasopressor", "vaso"],
        ),
        "rrt": _contains_any(supports_text, ["rrt", "hd", "sled", "dialysis", "cvvh"]),
        "o2": _contains_any(supports_text, ["o2", "oxygen", "hfnc", "nrbm"]),
    }


def _compute_flags(row: dict[str, str], status: str, support_flags: dict[str, bool]) -> tuple[list[str], str]:
    if status == "DECEASED":
        return ["DECEASED"], "RED"

    all_text = " ".join(str(value) for value in row.values()).lower()
    pending_text = row.get("pending", "").lower()
    new_issues_text = row.get("new_issues", "").lower()

    flags: list[str] = []
    if "dama" in all_text:
        flags.append("DAMA risk")
    if _contains_any(pending_text, ["ctpa", "ct brain", "ct head", "mri", "hrct", "ugie"]):
        flags.append("Pending critical imaging")
    if _contains_any(pending_text, ["neuro", "gastro", "cardiology", "cardio", "nephrology", "nephro", "ortho"]):
        flags.append("Pending consult")
    if support_flags.get("mv", False) and "high peak pressure" in new_issues_text:
        flags.append("Vent + high peak pressure")

    if any(flag in flags for flag in ("Pending critical imaging", "Vent + high peak pressure")):
        severity = "RED"
    elif flags:
        severity = "AMBER"
    else:
        severity = "NONE"
    return flags, severity


@lru_cache(maxsize=1)
def load_rules(rules_path: str = str(RULES_PATH)) -> dict[str, Any]:
    if yaml is None:
        raise RuntimeError("PyYAML is required. Install with `pip install pyyaml`.")
    with Path(rules_path).open("r", encoding="utf-8") as handle:
        return yaml.safe_load(handle) or {}


def _parse_required_items(payload: dict[str, Any], category: str) -> list[RequiredItem]:
    items_raw = payload.get(f"required_{category}", []) or []
    items: list[RequiredItem] = []
    for raw in items_raw:
        if not isinstance(raw, dict):
            continue
        name = _normalize_text(str(raw.get("name", "")))
        if not name:
            continue
        priority = _normalize_text(str(raw.get("priority", "Medium"))) or "Medium"
        trigger_keywords = [str(x) for x in (raw.get("trigger_keywords", []) or [])]
        evidence_keywords = [str(x) for x in (raw.get("evidence_keywords", []) or [])]
        if not evidence_keywords:
            evidence_keywords = [name]
        items.append(
            RequiredItem(
                category=category,
                name=name,
                priority=priority,
                trigger_keywords=trigger_keywords,
                evidence_keywords=evidence_keywords,
            )
        )
    return items


@lru_cache(maxsize=1)
def load_algorithms(root_path: str = str(ALGORITHM_ROOT)) -> dict[str, list[AlgorithmSpec]]:
    if yaml is None:
        raise RuntimeError("PyYAML is required. Install with `pip install pyyaml`.")
    root = Path(root_path)
    grouped: dict[str, list[AlgorithmSpec]] = {}
    if not root.exists():
        return grouped

    for file_path in sorted(root.rglob("*.yaml")):
        with file_path.open("r", encoding="utf-8") as handle:
            payload = yaml.safe_load(handle) or {}

        system_tag = str(payload.get("system_tag", file_path.parent.name))
        required_items = []
        required_items.extend(_parse_required_items(payload, "tests"))
        required_items.extend(_parse_required_items(payload, "imaging"))
        required_items.extend(_parse_required_items(payload, "consults"))
        required_items.extend(_parse_required_items(payload, "care_checks"))

        algo = AlgorithmSpec(
            name=str(payload.get("name", file_path.stem)),
            system_tag=system_tag,
            diagnosis_keywords=[str(x) for x in (payload.get("diagnosis_keywords", []) or [])],
            supports_keywords=[str(x) for x in (payload.get("supports_keywords", []) or [])],
            lab_keywords=[str(x) for x in (payload.get("lab_keywords", []) or [])],
            strict_diagnosis_keywords=[str(x) for x in (payload.get("strict_diagnosis_keywords", []) or [])],
            required_items=required_items,
        )
        grouped.setdefault(system_tag, []).append(algo)
    return grouped


def _renal_gate(row: dict[str, str]) -> bool:
    diagnosis_text = _normalize_text(f"{row.get('diagnosis', '')} {row.get('new_issues', '')}").lower()
    labs_text = _normalize_text(row.get("key_labs_imaging", "")).lower()
    supports_text = _normalize_text(row.get("supports", "")).lower()
    if _contains_any(diagnosis_text, ["aki", "acute kidney injury", "ckd", "chronic kidney disease", "uremia"]):
        return True
    creat_match = re.search(r"creat(?:inine)?[^0-9]{0,10}(\d+(?:\.\d+)?)", labs_text)
    if creat_match and float(creat_match.group(1)) >= 2.0:
        return True
    if _contains_any(f"{supports_text} {labs_text}", ["hd", "sled", "dialysis", "rrt"]):
        return True
    return False


def _detect_system_tags(row: dict[str, str], rules: dict[str, Any], max_tags: int = 3) -> list[str]:
    diagnosis_text = _normalize_text(f"{row.get('diagnosis', '')} {row.get('new_issues', '')}").lower()
    supports_text = _normalize_text(row.get("supports", "")).lower()
    scores: dict[str, int] = {}

    for system_tag, config in (rules.get("system_detection", {}) or {}).items():
        keywords = [str(x) for x in ((config or {}).get("diagnosis_keywords", []) or [])]
        scores[str(system_tag)] = sum(1 for keyword in keywords if keyword.lower() in diagnosis_text)

    # Heuristic boosters for common mixed-system beds.
    if _contains_any(f"{diagnosis_text} {supports_text}", ["copd", "t2rf", "type 2 respiratory failure", "mv", "ventilator"]):
        scores["07_respiratory"] = scores.get("07_respiratory", 0) + 2
    if _contains_any(diagnosis_text, ["meningitis", "sepsis", "septic", "pneumonia", "infection", "fever"]):
        scores["10_infectious"] = scores.get("10_infectious", 0) + 1
    if _renal_gate(row):
        scores["02_renal"] = scores.get("02_renal", 0) + 1

    ranked = sorted(
        [(tag, score) for tag, score in scores.items() if score > 0],
        key=lambda item: (-item[1], item[0]),
    )
    tags = [tag for tag, _score in ranked[:max_tags]]

    if tags:
        return tags
    if _contains_any(supports_text, ["mv", "ventilator", "mechanical ventilation"]):
        return ["07_respiratory"]
    if _contains_any(supports_text, ["rrt", "hd", "sled", "dialysis"]):
        return ["02_renal"]
    return ["11_misc"]


def _supports_gate(algo: AlgorithmSpec, row: dict[str, str]) -> bool:
    supports_text = _normalize_text(row.get("supports", "")).lower()
    if _normalized_key(algo.name) == _normalized_key("ARDS_VentStrategy"):
        return _contains_any(supports_text, ["mv", "ventilator", "mechanical ventilation"])
    return True


def _score_algorithm(algo: AlgorithmSpec, row: dict[str, str]) -> int:
    diagnosis_text = _normalize_text(f"{row.get('diagnosis', '')} {row.get('new_issues', '')}").lower()
    supports_text = _normalize_text(row.get("supports", "")).lower()
    labs_text = _normalize_text(row.get("key_labs_imaging", "")).lower()

    if algo.strict_diagnosis_keywords:
        diag_match = any(_contains_strict_keyword(diagnosis_text, keyword) for keyword in algo.strict_diagnosis_keywords)
    else:
        diag_match = _contains_any(diagnosis_text, algo.diagnosis_keywords)

    supports_match = _contains_any(supports_text, algo.supports_keywords) if algo.supports_keywords else False
    lab_match = _contains_any(labs_text, algo.lab_keywords) if algo.lab_keywords else False

    score = 0
    if diag_match:
        score += 2
    if supports_match:
        score += 1
    if lab_match:
        score += 1
    return score


def _select_algorithms(
    row: dict[str, str],
    system_tags: list[str],
    grouped_algorithms: dict[str, list[AlgorithmSpec]],
) -> list[AlgorithmSpec]:
    candidates: list[AlgorithmSpec] = []
    for tag in system_tags:
        candidates.extend(grouped_algorithms.get(tag, []))

    unique_candidates: dict[str, AlgorithmSpec] = {algo.name: algo for algo in candidates}
    scored: list[tuple[AlgorithmSpec, int]] = []
    for algo in unique_candidates.values():
        if algo.system_tag == "02_renal" and not _renal_gate(row):
            continue
        if not _supports_gate(algo, row):
            continue
        score = _score_algorithm(algo, row)
        if score >= 2:
            scored.append((algo, score))

    scored.sort(key=lambda item: (-item[1], item[0].name.lower()))
    return [algo for algo, _score in scored[:3]]


def _gather_evidence_text(row: dict[str, str]) -> str:
    fields = [
        row.get("actions_done", ""),
        row.get("plan_next_12h", ""),
        row.get("pending", ""),
        row.get("key_labs_imaging", ""),
    ]
    return _normalize_text(" ".join(fields)).lower()


def _expand_keyword_variants(keywords: list[str]) -> list[str]:
    expanded: list[str] = []
    for keyword in keywords:
        expanded.append(keyword)
        keyword_key = _normalized_key(keyword)
        for group in SYNONYM_GROUPS:
            if any(_normalized_key(entry) == keyword_key for entry in group):
                expanded.extend(group)
    return _dedupe_strings([_normalize_text(x) for x in expanded if _normalize_text(x)])


def _item_triggered(item: RequiredItem, problem_text: str) -> bool:
    if not item.trigger_keywords:
        return True
    return _contains_any(problem_text, item.trigger_keywords)


def _item_present(item: RequiredItem, evidence_text: str) -> bool:
    terms = _expand_keyword_variants(item.evidence_keywords if item.evidence_keywords else [item.name])
    evidence_key = _normalized_key(evidence_text)
    for term in terms:
        term_lower = term.lower()
        if term_lower in evidence_text:
            return True
        term_key = _normalized_key(term_lower)
        if term_key and term_key in evidence_key:
            return True
    return False


def _dedupe_required(items: list[RequiredItem]) -> list[RequiredItem]:
    seen: set[tuple[str, str]] = set()
    out: list[RequiredItem] = []
    for item in items:
        key = (item.category, _normalized_key(item.name))
        if key in seen:
            continue
        seen.add(key)
        out.append(item)
    return out


def _format_priority_item(item: RequiredItem) -> str:
    return f"{item.priority.title()}: {item.name}"


def _format_category(items: list[RequiredItem], category: str) -> str:
    lines = [_format_priority_item(item) for item in items if item.category == category]
    if not lines:
        return ""
    return "\n".join(f"- {line}" for line in lines)


def _build_missing_engine(
    row: dict[str, str],
    matched_algorithms: list[AlgorithmSpec],
    dama_risk: bool,
) -> tuple[list[RequiredItem], list[RequiredItem], list[RequiredItem]]:
    problem_text = _normalize_text(
        f"{row.get('diagnosis', '')} {row.get('new_issues', '')} {row.get('supports', '')} {row.get('status', '')}"
    ).lower()
    evidence_text = _gather_evidence_text(row)

    required_union: list[RequiredItem] = []
    for algo in matched_algorithms:
        for item in algo.required_items:
            if _item_triggered(item, problem_text):
                required_union.append(item)
    required_union = _dedupe_required(required_union)

    present: list[RequiredItem] = []
    missing: list[RequiredItem] = []
    for item in required_union:
        if _item_present(item, evidence_text):
            present.append(item)
        else:
            missing.append(item)

    missing.sort(key=lambda item: (_priority_rank(item.priority), item.category, item.name.lower()))
    present.sort(key=lambda item: (_priority_rank(item.priority), item.category, item.name.lower()))

    missing_main = [item for item in missing if item.category in {"tests", "imaging", "consults"}]
    missing_care = [item for item in missing if item.category == "care_checks"]

    if dama_risk:
        missing_main = [item for item in missing_main if item.priority.lower() == "high"]
        missing_care = [item for item in missing_care if item.priority.lower() == "high"]

    return missing_main[:5], missing_care[:5], present


def _status_rank(status: str) -> int:
    group = _status_group(status)
    return {"CRITICAL": 0, "SICK": 1, "SERIOUS": 2, "DECEASED": 3}.get(group, 4)


def _extract_datetime_from_filename(name: str) -> datetime | None:
    lower_name = name.lower()
    patterns: list[tuple[str, str]] = [
        (r"(20\d{2})[-_]?([01]\d)[-_]?([0-3]\d)[-_]?([0-2]\d)([0-5]\d)", "%Y%m%d%H%M"),
        (r"(20\d{2})[-_]?([01]\d)[-_]?([0-3]\d)", "%Y%m%d"),
        (r"([0-3]\d)[-_]([01]\d)[-_](20\d{2})[-_]?([0-2]\d)([0-5]\d)", "%d%m%Y%H%M"),
        (r"([0-3]\d)[-_]([01]\d)[-_](20\d{2})", "%d%m%Y"),
    ]
    for pattern, fmt in patterns:
        match = re.search(pattern, lower_name)
        if not match:
            continue
        token = "".join(match.groups())
        try:
            return datetime.strptime(token, fmt)
        except ValueError:
            continue
    return None


def infer_round_file_order(file_names: list[str]) -> tuple[int, int, str]:
    """
    Return (older_index, newer_index, reason) based on filename date/time and keywords.
    Falls back to listed order when hints are insufficient.
    """
    if len(file_names) < 2:
        raise ValueError("Need at least two filenames to infer round order.")

    first_name = file_names[0].lower()
    second_name = file_names[1].lower()
    first_dt = _extract_datetime_from_filename(first_name)
    second_dt = _extract_datetime_from_filename(second_name)

    if first_dt and second_dt and first_dt != second_dt:
        if first_dt < second_dt:
            return 0, 1, "filename date/time"
        return 1, 0, "filename date/time"

    def _hint_score(name: str, hints: list[str]) -> int:
        return sum(1 for hint in hints if hint in name)

    first_old = _hint_score(first_name, OLDER_FILENAME_HINTS)
    first_new = _hint_score(first_name, NEWER_FILENAME_HINTS)
    second_old = _hint_score(second_name, OLDER_FILENAME_HINTS)
    second_new = _hint_score(second_name, NEWER_FILENAME_HINTS)

    if first_old > first_new and second_new > second_old:
        return 0, 1, "older/newer filename keywords"
    if second_old > second_new and first_new > first_old:
        return 1, 0, "older/newer filename keywords"

    return 0, 1, "upload order fallback"


def _trend_from_previous(previous: dict[str, Any] | None, current: dict[str, Any]) -> tuple[str, str, bool]:
    if previous is None:
        return "NEW ADMISSION", "", False

    previous_status = str(previous.get("_status_group", "OTHER"))
    current_status = str(current.get("_status_group", "OTHER"))
    previous_severity = _status_severity(previous_status)
    current_severity = _status_severity(current_status)

    reasons: list[str] = []
    deteriorated = False

    if current_severity > previous_severity:
        deteriorated = True
        reasons.append(f"Status {previous_status} -> {current_status}")

    previous_support_score = _support_burden_score(previous)
    current_support_score = _support_burden_score(current)
    if current_support_score > previous_support_score:
        deteriorated = True
        reasons.append("Support escalation")

    support_labels = [
        ("_is_mv", "MV started"),
        ("_is_niv", "NIV/BiPAP started"),
        ("_is_vaso", "Vasopressor started"),
        ("_is_rrt", "RRT/HD/SLED started"),
    ]
    for key, label in support_labels:
        if bool(current.get(key)) and not bool(previous.get(key)):
            deteriorated = True
            reasons.append(label)

    if deteriorated:
        return "DETERIORATED", " | ".join(_dedupe_strings(reasons)), True

    if current_severity < previous_severity or current_support_score < previous_support_score:
        return "IMPROVED", "", False
    return "STABLE", "", False


def compare_rounds_outputs(previous_rows: list[dict[str, Any]], current_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    previous_map: dict[str, dict[str, Any]] = {}
    for row in previous_rows:
        key = _patient_key(row)
        if key:
            previous_map[key] = row

    compared: list[dict[str, Any]] = []
    for row in current_rows:
        row_copy = dict(row)
        key = _patient_key(row_copy)
        previous = previous_map.get(key) if key else None
        trend, reason, is_deteriorated = _trend_from_previous(previous, row_copy)

        row_copy["Round trend"] = trend
        row_copy["Deterioration since last round"] = "YES" if is_deteriorated else "NO"
        row_copy["Deterioration reasons"] = reason
        row_copy["_is_deteriorated"] = is_deteriorated
        compared.append(row_copy)

    return compared


def build_all_beds_outputs(table_rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    rules = load_rules()
    algorithms = load_algorithms()
    outputs: list[dict[str, Any]] = []

    for row in table_rows:
        bed = _normalize_text(row.get("bed", ""))
        if not bed:
            continue

        combined_text = " ".join(str(value) for value in row.values()).lower()
        support_flags = _support_markers(row)
        raw_pending = _normalize_text(row.get("pending", ""))
        dama_risk = "dama" in combined_text
        system_tags = _detect_system_tags(row, rules, max_tags=3)

        if "declared clinically dead" in combined_text:
            status = "DECEASED"
            matched_algorithms: list[AlgorithmSpec] = []
            missing_main: list[RequiredItem] = []
            missing_care: list[RequiredItem] = []
            present: list[RequiredItem] = []
            pending_verbatim = "Pending admin: trace reports" if raw_pending else ""
        else:
            status = _normalize_text(row.get("status", ""))
            matched_algorithms = _select_algorithms(row, system_tags, algorithms)
            missing_main, missing_care, present = _build_missing_engine(row, matched_algorithms, dama_risk)
            pending_verbatim = row.get("pending", "")

        flags, severity = _compute_flags(row, _status_group(status), support_flags)
        matched_names = [algo.name for algo in matched_algorithms]

        outputs.append(
            {
                "Bed": bed,
                "Patient ID": _normalize_text(row.get("patient_id", "")),
                "Diagnosis": _normalize_text(row.get("diagnosis", "")),
                "Status": status,
                "Supports": _normalize_text(row.get("supports", "")),
                "System tag": " | ".join(system_tags),
                "Matched algorithms": ", ".join(matched_names),
                "Missing Tests": _format_category(missing_main, "tests"),
                "Missing Imaging": _format_category(missing_main, "imaging"),
                "Missing Consults": _format_category(missing_main, "consults"),
                "Care checks (deterministic)": _format_category(missing_care, "care_checks"),
                "Pending (verbatim)": pending_verbatim,
                "Key labs/imaging (1 line)": _normalize_text(row.get("key_labs_imaging", "")),
                "Round trend": "",
                "Deterioration since last round": "",
                "Deterioration reasons": "",
                "_status_group": _status_group(status),
                "_flags": flags,
                "_flag_severity": severity,
                "_is_mv": support_flags["mv"],
                "_is_niv": support_flags["niv"],
                "_is_vaso": support_flags["vaso"],
                "_is_rrt": support_flags["rrt"],
                "_is_o2": support_flags["o2"],
                "_pending_verbatim": pending_verbatim,
                "_raw_new_issues": _normalize_text(row.get("new_issues", "")),
                "_system_tags": system_tags,
                "_matched_algorithms": matched_names,
                "_covered_tests": _format_category(present, "tests"),
                "_covered_imaging": _format_category(present, "imaging"),
                "_covered_consults": _format_category(present, "consults"),
                "_covered_care_checks": _format_category(present, "care_checks"),
                "_dama_risk": dama_risk,
                "_is_deteriorated": False,
            }
        )
    return outputs


def to_csv_bytes(rows: list[dict[str, Any]]) -> bytes:
    if not rows:
        return b""
    buffer = io.StringIO()
    writer = csv.DictWriter(
        buffer,
        fieldnames=DISPLAY_COLUMNS,
        quoting=csv.QUOTE_ALL,
        lineterminator="\n",
    )
    writer.writeheader()
    for row in rows:
        export_row = {}
        for column in DISPLAY_COLUMNS:
            export_row[column] = str(row.get(column, "")).replace("\n", " | ")
        writer.writerow(export_row)
    return buffer.getvalue().encode("utf-8")


def to_docx_bytes(rows: list[dict[str, Any]]) -> bytes:
    if Document is None:
        return b""
    sorted_rows = sorted(rows, key=lambda row: (_status_rank(str(row.get("Status", ""))), str(row.get("Bed", ""))))
    doc = Document()
    doc.add_heading("ALL Beds Output", level=1)
    if not sorted_rows:
        doc.add_paragraph("No rows generated.")
    else:
        table = doc.add_table(rows=1, cols=len(DISPLAY_COLUMNS))
        table.style = "Table Grid"
        for idx, header in enumerate(DISPLAY_COLUMNS):
            table.rows[0].cells[idx].text = header
        for row in sorted_rows:
            cells = table.add_row().cells
            for idx, header in enumerate(DISPLAY_COLUMNS):
                cells[idx].text = str(row.get(header, "")).replace("\n", " | ")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
