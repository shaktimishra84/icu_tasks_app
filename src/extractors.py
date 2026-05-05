from __future__ import annotations

import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Callable

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency at runtime
    Document = None

try:
    from PIL import Image
except Exception:  # pragma: no cover - optional dependency at runtime
    Image = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency at runtime
    PdfReader = None

try:
    import pytesseract
except Exception:  # pragma: no cover - optional dependency at runtime
    pytesseract = None

LOGGER = logging.getLogger(__name__)


class ExtractionError(Exception):
    """Raised when text extraction fails."""


DOCX_REQUIRED_FIELDS = [
    "bed",
    "patient_id",
    "diagnosis",
    "status",
    "supports",
    "new_issues",
    "actions_done",
    "plan_next_12h",
    "pending",
    "key_labs_imaging",
]

DOCX_HEADER_SEQUENCE = [
    "bed",
    "patient_id",
    "diagnosis",
    "status",
    "supports",
    "new_issues",
    "actions_done",
    "plan_next_12h",
    "pending",
    "key_labs_imaging",
]

DOCX_HEADER_ALIASES = {
    "bed": ["bed", "icu bed", "bed no", "bed number", "bed #", "bed id"],
    "patient_id": ["patient id", "patientid", "mrn", "uhid", "id", "patient no"],
    "diagnosis": ["diagnosis", "dx", "working diagnosis"],
    "status": ["status", "condition", "clinical status"],
    "supports": ["supports", "support", "organ supports", "devices", "device support"],
    "new_issues": ["new issues", "new issue", "issues", "new problems", "problem updates"],
    "actions_done": ["actions done", "done", "actions", "interventions done", "treatment done"],
    "plan_next_12h": ["plan next 12h", "plan next 12 h", "next 12h plan", "next 12 h plan", "plan 12h", "plan"],
    "pending": ["pending", "pending items", "pending tests", "awaiting", "to follow", "pending workup"],
    "key_labs_imaging": [
        "key labs imaging",
        "key labs/imaging",
        "labs/imaging",
        "key investigations",
        "investigations",
        "labs and imaging",
        "key labs",
    ],
}

DOCX_TEMPLATE_TOKENS = {
    "(1 line)",
    "(2 lines)",
    "(3 lines)",
    "(stable/watch/sick/crit.)",
    "(stable/watch/sick/crit)",
    "(stable watch sick crit.)",
    "(stable watch sick crit)",
    "(yes/no)",
}


def _normalize_header(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text.strip().lower())
    normalized = normalized.replace("_", " ")
    normalized = normalized.replace("/", " ")
    normalized = normalized.replace("-", " ")
    normalized = re.sub(r"[^a-z0-9 #]+", "", normalized)
    return normalized


def _alias_score(cell_text: str, alias_text: str) -> int:
    cell = _normalize_header(cell_text)
    alias = _normalize_header(alias_text)
    if not cell or not alias:
        return 0
    if cell == alias:
        return 5
    if re.search(rf"\b{re.escape(alias)}\b", cell):
        return 4
    if alias in cell:
        return 3
    if len(cell) >= 5 and cell in alias:
        return 2

    cell_tokens = {token for token in cell.split() if token}
    alias_tokens = {token for token in alias.split() if token}
    if not cell_tokens or not alias_tokens:
        return 0
    overlap = len(cell_tokens & alias_tokens)
    if overlap >= 2:
        return 2
    if overlap == 1 and max(len(token) for token in alias_tokens) >= 4:
        return 1
    return 0


def _canonical_docx_field(header: str) -> str | None:
    best_field: str | None = None
    best_score = 0
    for canonical_key, aliases in DOCX_HEADER_ALIASES.items():
        for alias in aliases:
            score = _alias_score(header, alias)
            if score > best_score:
                best_field = canonical_key
                best_score = score
    if best_score >= 3:
        return best_field
    return None


def _pad_row(cells: list[str], size: int) -> list[str]:
    return cells + [""] * max(0, size - len(cells))


def _extract_bed_digits(value: str) -> str:
    text = str(value or "").strip()
    bed_match = re.search(r"\bBED\s*([A-Za-z]?\d+[A-Za-z]?)\b", text, flags=re.IGNORECASE)
    if bed_match:
        return bed_match.group(1).upper()
    matches = re.findall(r"\d+[A-Za-z]?", text)
    return matches[-1].upper() if matches else ""


def _is_template_value(value: str) -> bool:
    normalized = re.sub(r"\s+", " ", value.strip().lower())
    if not normalized:
        return True
    if normalized in DOCX_TEMPLATE_TOKENS:
        return True
    if normalized.startswith("(") and normalized.endswith(")") and "line" in normalized:
        return True
    if "stable/watch/sick/crit" in normalized:
        return True
    return False


def _row_matches_expected_header(cells: list[str]) -> bool:
    return _match_header_columns(cells) is not None


def _match_header_columns(cells: list[str]) -> dict[str, int] | None:
    if not cells:
        return None

    candidates: list[tuple[int, int, str]] = []
    for idx, cell in enumerate(cells):
        for field in DOCX_HEADER_SEQUENCE:
            for alias in DOCX_HEADER_ALIASES.get(field, []):
                score = _alias_score(cell, alias)
                if score > 0:
                    candidates.append((score, idx, field))

    if not candidates:
        return None

    candidates.sort(key=lambda item: (-item[0], item[1], DOCX_HEADER_SEQUENCE.index(item[2])))
    field_map: dict[str, int] = {}
    used_idx: set[int] = set()
    for score, idx, field in candidates:
        if idx in used_idx:
            continue
        if field in field_map:
            continue
        # weak token overlaps are too noisy for header detection.
        if score < 2:
            continue
        field_map[field] = idx
        used_idx.add(idx)

    required = {"bed", "patient_id", "diagnosis"}
    if not required.issubset(field_map):
        return None
    if len(field_map) < 6:
        return None
    return field_map


def _is_header_like_row(cells: list[str], header_map: dict[str, int]) -> bool:
    row_map = _match_header_columns(cells)
    if row_map is None:
        return False
    overlap = set(row_map) & set(header_map)
    return len(overlap) >= 6


def _cell_at(cells: list[str], index: int | None) -> str:
    if index is None:
        return ""
    if index < 0 or index >= len(cells):
        return ""
    return str(cells[index]).strip()


def _extract_row_values(cells: list[str], header_map: dict[str, int]) -> dict[str, str]:
    row_values: dict[str, str] = {}
    for field in DOCX_HEADER_SEQUENCE[1:]:
        value = _cell_at(cells, header_map.get(field))
        if value and not _is_template_value(value):
            row_values[field] = value
    return row_values


def _patient_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())


def _is_plausible_patient_id(value: str) -> bool:
    raw = str(value or "").strip()
    if any(char in raw for char in ["/", ",", ".", "(", ")"]):
        return False
    if re.search(r"\b\d{1,2}(?:[:.]\d{2})?\s*(?:AM|PM)\b", raw, flags=re.IGNORECASE):
        return False
    if re.search(r"\b\d{2,3}\s*/\s*\d{2,3}(?:\s*\(\s*\d{2,3}\s*\))?", raw):
        return False
    if re.search(r"\bE\d+\s*V(?:T|NT|\d+)\s*M\d+\b", raw, flags=re.IGNORECASE):
        return False
    compact_raw = re.sub(r"[^A-Z0-9]+", "", raw.upper())
    if re.fullmatch(r"E\d+V[A-Z0-9]*M\d+", compact_raw):
        return False
    if re.search(r"\bDR\.?\b", raw, flags=re.IGNORECASE):
        return False
    cleaned = re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())
    if not cleaned:
        return False
    if cleaned.isdigit():
        return len(cleaned) >= 8
    if any(char.isalpha() for char in cleaned):
        if not any(char.isdigit() for char in cleaned):
            return False
        if re.search(r"\s", raw):
            return False
        return len(cleaned) >= 4
    return len(cleaned) >= 6


def _append_record_value(record: dict[str, str], field: str, value: str) -> None:
    cleaned = value.strip()
    if not cleaned or _is_template_value(cleaned):
        return
    existing = str(record.get(field, "")).strip()
    if not existing:
        record[field] = cleaned
        return
    existing_parts = [part.strip().lower() for part in existing.split("|")]
    if cleaned.lower() in existing_parts:
        return
    record[field] = f"{existing} | {cleaned}"


def _dedupe_parsed_rows(parsed_rows: list[dict[str, str]]) -> tuple[list[dict[str, str]], list[str]]:
    warnings: list[str] = []
    deduped_rows: list[dict[str, str]] = []
    seen_patient_ids: dict[str, int] = {}
    for row in parsed_rows:
        patient_id = str(row.get("patient_id", "")).strip()
        patient_key = _patient_key(patient_id)
        if not patient_key:
            deduped_rows.append(row)
            continue

        if patient_key not in seen_patient_ids:
            seen_patient_ids[patient_key] = len(deduped_rows)
            deduped_rows.append(row)
            continue

        existing_index = seen_patient_ids[patient_key]
        existing = deduped_rows[existing_index]
        existing_bed = str(existing.get("bed", "")).strip()
        row_bed = str(row.get("bed", "")).strip()

        if existing_bed != row_bed:
            warnings.append(
                f"Duplicate patient ID `{patient_id}` found in beds {existing_bed} and {row_bed}; "
                f"kept first occurrence (bed {existing_bed})."
            )
            continue

        for field in DOCX_REQUIRED_FIELDS[1:]:
            _append_record_value(existing, field, str(row.get(field, "")))

    return deduped_rows, warnings


def _parse_rows_with_header_map(
    rows_as_cells: list[list[str]],
    header_map: dict[str, int],
    start_index: int = 0,
) -> tuple[list[dict[str, str]], list[str]]:
    parsed_rows: list[dict[str, str]] = []
    current_record: dict[str, str] | None = None

    for cells in rows_as_cells[start_index:]:
        if _is_header_like_row(cells, header_map):
            continue

        bed_cell = _cell_at(cells, header_map.get("bed"))
        bed_digits = _extract_bed_digits(bed_cell)
        row_values = _extract_row_values(cells, header_map)

        if not bed_digits and not row_values:
            continue

        has_patient_id = bool(row_values.get("patient_id", "").strip())
        has_core_fields = any(row_values.get(field, "").strip() for field in ("diagnosis", "status", "supports"))

        if bed_digits:
            if current_record is None:
                current_record = {field: "" for field in DOCX_REQUIRED_FIELDS}
                current_record["bed"] = bed_digits
                for field, value in row_values.items():
                    _append_record_value(current_record, field, value)
                continue

            current_bed = str(current_record.get("bed", "")).strip()
            if bed_digits == current_bed and not has_patient_id:
                for field, value in row_values.items():
                    _append_record_value(current_record, field, value)
                continue

            if bed_digits != current_bed and not has_patient_id and not has_core_fields:
                for field, value in row_values.items():
                    _append_record_value(current_record, field, value)
                continue

            if current_record.get("bed"):
                parsed_rows.append(current_record)

            current_record = {field: "" for field in DOCX_REQUIRED_FIELDS}
            current_record["bed"] = bed_digits
            for field, value in row_values.items():
                _append_record_value(current_record, field, value)
            continue

        if current_record is None:
            continue
        for field, value in row_values.items():
            _append_record_value(current_record, field, value)

    if current_record is not None and current_record.get("bed"):
        parsed_rows.append(current_record)

    if not parsed_rows:
        return [], []
    return _dedupe_parsed_rows(parsed_rows)


def _fallback_header_maps(rows_as_cells: list[list[str]]) -> list[dict[str, int]]:
    if not rows_as_cells:
        return []
    max_columns = max(len(row) for row in rows_as_cells)
    maps: list[dict[str, int]] = []
    for offset in (0, 1, 2):
        field_map: dict[str, int] = {}
        for idx, field in enumerate(DOCX_HEADER_SEQUENCE):
            column_index = idx + offset
            if column_index < max_columns:
                field_map[field] = column_index
        if {"bed", "patient_id", "diagnosis"}.issubset(field_map):
            maps.append(field_map)
    return maps


def _is_structured_rmo_section_table(rows_as_cells: list[list[str]]) -> bool:
    first_rows = " ".join(" ".join(row) for row in rows_as_cells[:2])
    header = _normalize_header(first_rows)
    section_markers = [
        ("uhid" in header and "patient name" in header and "icu bed no" in header),
        ("gcs" in header and "bp map" in header),
        ("mode" in header and ("fio2" in header or "peep" in header)),
        ("map stability" in header or "inotropes" in header),
        ("gcs trend" in header),
        ("urine trend" in header and "creatinine" in header),
        ("red flag present" in header),
        ("current status" in header and "major concern" in header),
        ("advice given" in header),
    ]
    return any(section_markers)


def _parse_bed_table(rows_as_cells: list[list[str]]) -> tuple[list[dict[str, str]], list[str]]:
    warnings: list[str] = []
    best_header_index: int | None = None
    best_header_map: dict[str, int] | None = None
    best_score = -1
    for index, cells in enumerate(rows_as_cells):
        header_map = _match_header_columns(cells)
        if header_map is None:
            continue
        score = len(header_map)
        if score > best_score:
            best_score = score
            best_header_index = index
            best_header_map = header_map

    if best_header_index is not None and best_header_map is not None:
        parsed_rows, warnings = _parse_rows_with_header_map(
            rows_as_cells=rows_as_cells,
            header_map=best_header_map,
            start_index=best_header_index + 1,
        )
        if parsed_rows:
            return parsed_rows, warnings

    if _is_structured_rmo_section_table(rows_as_cells):
        return [], warnings

    fallback_best_rows: list[dict[str, str]] = []
    fallback_best_warnings: list[str] = []
    fallback_best_score = -1
    fallback_best_plausible_pid = 0
    for fallback_map in _fallback_header_maps(rows_as_cells):
        parsed_rows, fallback_warnings = _parse_rows_with_header_map(
            rows_as_cells=rows_as_cells,
            header_map=fallback_map,
            start_index=0,
        )
        if not parsed_rows:
            continue
        plausible_pid = sum(
            1 for row in parsed_rows if _is_plausible_patient_id(str(row.get("patient_id", "")))
        )
        if plausible_pid <= 0:
            continue
        any_pid = sum(1 for row in parsed_rows if str(row.get("patient_id", "")).strip())
        score = (plausible_pid * 1000) + (any_pid * 20) + len(parsed_rows)
        if score > fallback_best_score:
            fallback_best_score = score
            fallback_best_plausible_pid = plausible_pid
            fallback_best_rows = parsed_rows
            fallback_best_warnings = fallback_warnings

    if fallback_best_rows and fallback_best_plausible_pid > 0:
        warnings = list(fallback_best_warnings)
        warnings.append("Header row not confidently detected; used positional fallback parsing.")
        return fallback_best_rows, warnings

    return [], warnings


def _select_best_table_rows(
    table_rows_collection: list[list[list[str]]],
) -> tuple[list[dict[str, str]], int | None, list[str]]:
    best_rows: list[dict[str, str]] = []
    best_table_index: int | None = None
    best_score = -1
    best_warnings: list[str] = []

    for table_index, rows_as_cells in enumerate(table_rows_collection):
        parsed_rows, warnings = _parse_bed_table(rows_as_cells)
        if not parsed_rows:
            continue

        plausible_pid = sum(
            1 for row in parsed_rows if _is_plausible_patient_id(str(row.get("patient_id", "")))
        )
        any_pid = sum(1 for row in parsed_rows if str(row.get("patient_id", "")).strip())
        score = (plausible_pid * 1000) + (any_pid * 20) + len(parsed_rows)
        if score > best_score:
            best_score = score
            best_rows = parsed_rows
            best_table_index = table_index
            best_warnings = warnings

    return best_rows, best_table_index, best_warnings


def _extract_pdf(data: bytes) -> str:
    if PdfReader is None:
        raise ExtractionError("PDF support missing. Install pypdf.")
    reader = PdfReader(BytesIO(data))
    text_parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(text_parts).strip()


def extract_pdf_pages(data: bytes) -> list[tuple[int, str]]:
    if PdfReader is None:
        raise ExtractionError("PDF support missing. Install pypdf.")
    reader = PdfReader(BytesIO(data))
    pages: list[tuple[int, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((index, text))
    return pages


def _extract_docx(data: bytes) -> str:
    payload = _extract_docx_content(data)
    return payload["raw_text"]


def _extract_docx_content(data: bytes) -> dict[str, Any]:
    if Document is None:
        raise ExtractionError("DOCX support missing. Install python-docx.")
    document = Document(BytesIO(data))

    paragraph_parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []

    table_count = len(document.tables)
    LOGGER.info("DOCX tables detected: %d", table_count)

    debug_raw_rows: list[list[str]] = []
    table_rows_collection: list[list[list[str]]] = []

    for table in document.tables:
        rows_as_cells: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_as_cells.append(cells)
            if any(cells):
                table_lines.append(" | ".join(cells))
                if len(debug_raw_rows) < 200:
                    debug_raw_rows.append(cells)
        table_rows_collection.append(rows_as_cells)

    parsed_rows, selected_table_index, parse_warnings = _select_best_table_rows(table_rows_collection)

    if paragraph_parts and table_lines:
        raw_text = "\n".join(paragraph_parts + [""] + table_lines).strip()
    elif paragraph_parts:
        raw_text = "\n".join(paragraph_parts).strip()
    else:
        raw_text = "\n".join(table_lines).strip()

    if selected_table_index is None:
        LOGGER.info("DOCX bed-table header not found in any table.")
    else:
        LOGGER.info("DOCX selected table index for bed parsing: %d", selected_table_index)
    LOGGER.info("DOCX table rows parsed: %d", len(parsed_rows))
    return {
        "raw_text": raw_text,
        "table_rows": parsed_rows,
        "debug_raw_rows": debug_raw_rows[:50],
        "table_index": selected_table_index,
        "parse_warnings": parse_warnings,
    }


def extract_docx_table_rows(data: bytes) -> list[dict[str, str]]:
    payload = _extract_docx_content(data)
    rows = payload.get("table_rows", [])
    return rows if isinstance(rows, list) else []


def _extract_text(data: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore").strip()


def _extract_image(data: bytes, image_ocr_fn: Callable[[bytes], str] | None = None) -> str:
    if Image is None:
        raise ExtractionError("Image support missing. Install Pillow.")
    image = Image.open(BytesIO(data))

    text = ""
    if pytesseract is not None:
        text = (pytesseract.image_to_string(image) or "").strip()

    if not text and image_ocr_fn is not None:
        text = (image_ocr_fn(data) or "").strip()

    if not text:
        raise ExtractionError(
            "Image OCR could not run. Install Tesseract locally or provide an OCR callback."
        )
    return text


def extract_text(
    filename: str,
    data: bytes,
    image_ocr_fn: Callable[[bytes], str] | None = None,
) -> str | dict[str, Any]:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx_content(data)
    elif ext in {".txt", ".md"}:
        text = _extract_text(data)
    elif ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        text = _extract_image(data, image_ocr_fn=image_ocr_fn)
    elif ext == ".doc":
        raise ExtractionError("Legacy .doc files are not supported. Convert to .docx first.")
    else:
        raise ExtractionError(f"Unsupported file type: {ext or '(no extension)'}")

    if isinstance(text, dict):
        raw_text = text.get("raw_text", "")
        table_rows = text.get("table_rows", [])
        if not raw_text and not table_rows:
            raise ExtractionError("No text could be extracted from this file.")
    elif not text:
        raise ExtractionError("No text could be extracted from this file.")
    return text
