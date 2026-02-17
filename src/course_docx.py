from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from html import escape as _xml_escape
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document
except Exception:  # pragma: no cover - optional at runtime
    Document = None

from src.rounds_tracker import compute_snapshot_changes, split_field_items, status_group_from_text, support_labels_from_state


def _safe_text(value: Any) -> str:
    clean = str(value or "")
    clean = clean.encode("utf-8", "replace").decode("utf-8")
    clean = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", " ", clean)
    return re.sub(r"\s+", " ", clean.strip())


def _round_label(row: dict[str, Any]) -> str:
    return f"{_safe_text(row.get('date', '-'))} ({_safe_text(row.get('shift', '-'))})"


def _pending_summary(chronological_rows: list[dict[str, Any]]) -> tuple[list[str], list[str]]:
    if not chronological_rows:
        return [], []

    latest_pending = split_field_items(chronological_rows[-1].get("pending", ""))
    latest_keys = {item.lower() for item in latest_pending}

    previous_items: dict[str, str] = {}
    for row in chronological_rows[:-1]:
        for item in split_field_items(row.get("pending", "")):
            previous_items.setdefault(item.lower(), item)

    resolved = [value for key, value in previous_items.items() if key not in latest_keys]
    return latest_pending, resolved


def _turning_points(chronological_rows: list[dict[str, Any]]) -> list[str]:
    out: list[str] = []
    for index in range(1, len(chronological_rows)):
        previous = chronological_rows[index - 1]
        current = chronological_rows[index]
        changes = compute_snapshot_changes([current], [previous])
        if not changes:
            continue
        change = changes[0]
        trend = str(change.get("trend", ""))
        if trend in {"DETERIORATED", "IMPROVED"} or change.get("supports_added") or change.get("supports_removed"):
            summary = ", ".join(change.get("summary_lines", [])[:2])
            out.append(f"{_round_label(current)}: {summary}")
    return out


def _add_heading(doc: Any, title: str, level: int = 2) -> None:
    doc.add_heading(title, level=level)


def _add_bullets(doc: Any, items: list[str], fallback: str = "-") -> None:
    if not items:
        doc.add_paragraph(fallback, style="List Bullet")
        return
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _sbar_blocks(
    chronological_rows: list[dict[str, Any]],
    selected_output: dict[str, Any] | None,
) -> dict[str, str]:
    admission = chronological_rows[0]
    current = chronological_rows[-1]
    admission_status = status_group_from_text(admission.get("status", ""))
    current_status = status_group_from_text(current.get("status", ""))

    situation = (
        f"Current round {_round_label(current)} with status {current_status}. "
        f"Primary diagnosis context: {_safe_text(current.get('diagnosis', '-'))}."
    )
    background = (
        f"Admission baseline {_round_label(admission)} status {admission_status}. "
        f"Total rounds documented: {len(chronological_rows)}. "
        f"Key trajectory notes: {('; '.join(_turning_points(chronological_rows)[:3]) or 'No major turning point flagged.')}"
    )
    assessment = (
        f"Supports currently active: {', '.join(support_labels_from_state(current)) or '-'}; "
        f"New issues: {'; '.join(split_field_items(current.get('new_issues', ''))[:4]) or '-'}; "
        f"Key labs/imaging: {_safe_text(current.get('key_labs_imaging', '-'))}."
    )

    recommendation_parts: list[str] = []
    plan = split_field_items(current.get("plan_next_12h", ""))
    pending = split_field_items(current.get("pending", ""))
    if plan:
        recommendation_parts.append("Plan: " + "; ".join(plan[:6]))
    if pending:
        recommendation_parts.append("Pending: " + "; ".join(pending[:6]))
    if selected_output:
        misses = []
        for key in ["Missing Tests", "Missing Imaging", "Missing Consults", "Care checks (deterministic)"]:
            misses.extend(split_field_items(selected_output.get(key, "")))
        if misses:
            recommendation_parts.append("Deterministic missed-item prompts: " + "; ".join(misses[:8]))
    recommendation = " ".join(recommendation_parts) or "No explicit recommendation recorded in source rounds."

    return {
        "Situation": situation,
        "Background": background,
        "Assessment": assessment,
        "Recommendation": recommendation,
    }


def _fallback_lines(
    *,
    unit_name: str,
    patient_label: str,
    chronological_rows: list[dict[str, Any]],
    selected_output: dict[str, Any] | None,
    resource_matches: list[tuple[Any, float]],
) -> list[str]:
    admission = chronological_rows[0]
    current = chronological_rows[-1]
    sbar = _sbar_blocks(chronological_rows, selected_output)
    lines: list[str] = [
        "ICU Course History and Handoff Summary",
        f"Patient: {patient_label}",
        f"Unit: {unit_name}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "1. Structured SBAR Handoff",
        f"Situation: {sbar['Situation']}",
        f"Background: {sbar['Background']}",
        f"Assessment: {sbar['Assessment']}",
        f"Recommendation: {sbar['Recommendation']}",
        "",
        "2. Course Since Admission",
        f"Admission baseline: {_round_label(admission)} | Status: {status_group_from_text(admission.get('status', ''))}",
        f"Current round: {_round_label(current)} | Status: {status_group_from_text(current.get('status', ''))}",
        "Major Turning Points:",
    ]
    points = _turning_points(chronological_rows)[:10]
    lines.extend([f"- {item}" for item in points] if points else ["- -"])

    support_lines: list[str] = []
    for label in ["MV", "NIV", "VASO", "RRT", "O2"]:
        start = -1
        end = -1
        for idx, row in enumerate(chronological_rows):
            if label in support_labels_from_state(row):
                if start < 0:
                    start = idx
                end = idx
        if start < 0:
            continue
        start_label = _round_label(chronological_rows[start])
        if end == len(chronological_rows) - 1:
            support_lines.append(f"{label}: started {start_label}, ongoing")
        else:
            end_label = _round_label(chronological_rows[end])
            support_lines.append(f"{label}: started {start_label}, stopped by {end_label}")
    lines.append("Support Trajectory:")
    lines.extend([f"- {item}" for item in support_lines] if support_lines else ["- -"])

    lines.append("")
    lines.append("3. Pending and Missed-Item Crosscheck")
    pending_now, resolved_pending = _pending_summary(chronological_rows)
    lines.append("Pending now:")
    lines.extend([f"- {item}" for item in pending_now] if pending_now else ["- -"])
    lines.append("Resolved since admission:")
    lines.extend([f"- {item}" for item in resolved_pending[:12]] if resolved_pending else ["- -"])

    if selected_output:
        sections = [
            ("Missing Tests", "Missing Tests"),
            ("Missing Imaging", "Missing Imaging"),
            ("Missing Consults", "Missing Consults"),
            ("Care Checks", "Care checks (deterministic)"),
        ]
        for title, field in sections:
            lines.append(title + ":")
            items = split_field_items(selected_output.get(field, ""))[:12]
            lines.extend([f"- {item}" for item in items] if items else ["- -"])

    lines.append("")
    lines.append("4. Resource Crosscheck Context")
    if resource_matches:
        for chunk, score in resource_matches[:8]:
            file_name = _safe_text(getattr(chunk, "file_name", "resource"))
            page = int(getattr(chunk, "page_number", 0) or 0)
            excerpt = _safe_text(getattr(chunk, "text", ""))[:200]
            lines.append(f"- {file_name} (p.{page}, score {score:.3f}) - {excerpt}")
    else:
        lines.append("- No indexed resource context found for this patient course.")

    lines.append("")
    lines.append("5. Professional Documentation Checklist")
    for item in [
        "Strategic handoff orientation: summarize what must happen next, not only what happened.",
        "Systems-based synthesis: respiratory, cardiovascular, neuro, renal, infection, and rehabilitation trajectory.",
        "High-fidelity antimicrobial and support data: indication, duration, escalation/de-escalation rationale.",
        "Code-status and goals-of-care clarity where applicable.",
        "Pending and unresolved tasks listed with explicit ownership at transfer.",
        "Patient-facing plain-language explanation should be prepared separately before transfer/discharge.",
    ]:
        lines.append(f"- {item}")
    return lines


def _minimal_docx_bytes(paragraphs: list[str]) -> bytes:
    paragraph_xml = []
    for text in paragraphs:
        clean = _safe_text(text)
        escaped = _xml_escape(clean)
        paragraph_xml.append(
            f"<w:p><w:r><w:t xml:space=\"preserve\">{escaped}</w:t></w:r></w:p>"
        )
    body_xml = "".join(paragraph_xml) + (
        "<w:sectPr>"
        "<w:pgSz w:w=\"12240\" w:h=\"15840\"/>"
        "<w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" "
        "w:header=\"708\" w:footer=\"708\" w:gutter=\"0\"/>"
        "</w:sectPr>"
    )
    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        f"<w:body>{body_xml}</w:body>"
        "</w:document>"
    )
    content_types_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
        "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
        "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
        "<Override PartName=\"/word/document.xml\" "
        "ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
        "<Override PartName=\"/docProps/core.xml\" "
        "ContentType=\"application/vnd.openxmlformats-package.core-properties+xml\"/>"
        "</Types>"
    )
    rels_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
        "<Relationship Id=\"rId1\" "
        "Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" "
        "Target=\"word/document.xml\"/>"
        "<Relationship Id=\"rId2\" "
        "Type=\"http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties\" "
        "Target=\"docProps/core.xml\"/>"
        "</Relationships>"
    )
    document_rels_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\"></Relationships>"
    )
    generated_ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    core_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<cp:coreProperties "
        "xmlns:cp=\"http://schemas.openxmlformats.org/package/2006/metadata/core-properties\" "
        "xmlns:dc=\"http://purl.org/dc/elements/1.1/\" "
        "xmlns:dcterms=\"http://purl.org/dc/terms/\" "
        "xmlns:dcmitype=\"http://purl.org/dc/dcmitype/\" "
        "xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\">"
        "<dc:title>ICU Course History and Handoff Summary</dc:title>"
        "<dc:creator>ICU Task Assistant</dc:creator>"
        f"<dcterms:created xsi:type=\"dcterms:W3CDTF\">{generated_ts}</dcterms:created>"
        f"<dcterms:modified xsi:type=\"dcterms:W3CDTF\">{generated_ts}</dcterms:modified>"
        "</cp:coreProperties>"
    )

    out = io.BytesIO()
    with ZipFile(out, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml)
        archive.writestr("_rels/.rels", rels_xml)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", document_rels_xml)
        archive.writestr("docProps/core.xml", core_xml)
    return out.getvalue()


def _docx_filename(patient_label: str) -> str:
    label_safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", patient_label).strip("_") or "patient"
    return f"ICU_Course_{label_safe}.docx"


def generate_course_docx(
    *,
    unit_name: str,
    patient_label: str,
    patient_key: str,
    chronological_rows: list[dict[str, Any]],
    selected_output: dict[str, Any] | None,
    resource_matches: list[tuple[Any, float]],
) -> tuple[bytes, str]:
    if not chronological_rows:
        raise ValueError("No course history rows found.")

    filename = _docx_filename(patient_label)

    if Document is None:
        fallback = _fallback_lines(
            unit_name=unit_name,
            patient_label=patient_label,
            chronological_rows=chronological_rows,
            selected_output=selected_output,
            resource_matches=resource_matches,
        )
        return _minimal_docx_bytes(fallback), filename

    doc = Document()
    doc.add_heading("ICU Course History and Handoff Summary", level=1)
    doc.add_paragraph(f"Patient: {patient_label}")
    doc.add_paragraph(f"Unit: {unit_name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    _add_heading(doc, "1. Structured SBAR Handoff", level=2)
    sbar = _sbar_blocks(chronological_rows, selected_output)
    for key in ["Situation", "Background", "Assessment", "Recommendation"]:
        doc.add_paragraph(f"{key}: {sbar[key]}")

    _add_heading(doc, "2. Course Since Admission", level=2)
    admission = chronological_rows[0]
    current = chronological_rows[-1]
    doc.add_paragraph(
        f"Admission baseline: {_round_label(admission)} | Status: {status_group_from_text(admission.get('status', ''))}"
    )
    doc.add_paragraph(
        f"Current round: {_round_label(current)} | Status: {status_group_from_text(current.get('status', ''))}"
    )

    _add_heading(doc, "Major Turning Points", level=3)
    _add_bullets(doc, _turning_points(chronological_rows)[:10])

    _add_heading(doc, "Support Trajectory", level=3)
    support_lines: list[str] = []
    for label in ["MV", "NIV", "VASO", "RRT", "O2"]:
        start = -1
        end = -1
        for idx, row in enumerate(chronological_rows):
            if label in support_labels_from_state(row):
                if start < 0:
                    start = idx
                end = idx
        if start < 0:
            continue
        start_label = _round_label(chronological_rows[start])
        if end == len(chronological_rows) - 1:
            support_lines.append(f"{label}: started {start_label}, ongoing")
        else:
            end_label = _round_label(chronological_rows[end])
            support_lines.append(f"{label}: started {start_label}, stopped by {end_label}")
    _add_bullets(doc, support_lines)

    _add_heading(doc, "3. Pending and Missed-Item Crosscheck", level=2)
    pending_now, resolved_pending = _pending_summary(chronological_rows)
    doc.add_paragraph("Pending now:")
    _add_bullets(doc, pending_now)
    doc.add_paragraph("Resolved since admission:")
    _add_bullets(doc, resolved_pending[:12])

    if selected_output:
        sections = [
            ("Missing Tests", "Missing Tests"),
            ("Missing Imaging", "Missing Imaging"),
            ("Missing Consults", "Missing Consults"),
            ("Care Checks", "Care checks (deterministic)"),
        ]
        for title, field in sections:
            _add_heading(doc, title, level=3)
            _add_bullets(doc, split_field_items(selected_output.get(field, ""))[:12])

    _add_heading(doc, "4. Resource Crosscheck Context", level=2)
    if resource_matches:
        for chunk, score in resource_matches[:8]:
            file_name = _safe_text(getattr(chunk, "file_name", "resource"))
            page = int(getattr(chunk, "page_number", 0) or 0)
            excerpt = _safe_text(getattr(chunk, "text", ""))[:200]
            doc.add_paragraph(
                f"{file_name} (p.{page}, score {score:.3f}) - {excerpt}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No indexed resource context found for this patient course.", style="List Bullet")

    _add_heading(doc, "5. Professional Documentation Checklist", level=2)
    checklist = [
        "Strategic handoff orientation: summarize what must happen next, not only what happened.",
        "Systems-based synthesis: respiratory, cardiovascular, neuro, renal, infection, and rehabilitation trajectory.",
        "High-fidelity antimicrobial and support data: indication, duration, escalation/de-escalation rationale.",
        "Code-status and goals-of-care clarity where applicable.",
        "Pending and unresolved tasks listed with explicit ownership at transfer.",
        "Patient-facing plain-language explanation should be prepared separately before transfer/discharge.",
    ]
    _add_bullets(doc, checklist)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), filename


def generate_course_docx_safe(
    *,
    unit_name: str,
    patient_label: str,
    patient_key: str,
    chronological_rows: list[dict[str, Any]],
    selected_output: dict[str, Any] | None,
    resource_matches: list[tuple[Any, float]],
) -> tuple[bytes, str, str | None]:
    filename = _docx_filename(patient_label)
    try:
        docx_bytes, generated_name = generate_course_docx(
            unit_name=unit_name,
            patient_label=patient_label,
            patient_key=patient_key,
            chronological_rows=chronological_rows,
            selected_output=selected_output,
            resource_matches=resource_matches,
        )
        return docx_bytes, generated_name, None
    except Exception as error:
        warning = _safe_text(str(error)) or error.__class__.__name__
        safe_rows = chronological_rows
        if not safe_rows:
            safe_rows = [
                {
                    "date": datetime.now().strftime("%Y-%m-%d"),
                    "shift": "-",
                    "status": "",
                    "diagnosis": "",
                    "supports": "",
                    "new_issues": "",
                    "actions_done": "",
                    "plan_next_12h": "",
                    "pending": "",
                    "key_labs_imaging": "",
                }
            ]

        try:
            lines = _fallback_lines(
                unit_name=unit_name,
                patient_label=patient_label,
                chronological_rows=safe_rows,
                selected_output=selected_output,
                resource_matches=resource_matches,
            )
        except Exception:
            lines = [
                "ICU Course History and Handoff Summary",
                f"Patient: {patient_label}",
                f"Unit: {unit_name}",
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                "",
                "Fallback export generated due to an internal rendering error.",
            ]

        lines.insert(0, f"Export fallback note: {warning}")
        return _minimal_docx_bytes(lines), filename, warning
